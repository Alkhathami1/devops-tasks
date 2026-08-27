#!/usr/bin/env python3
"""
Render docs/REPORT.md to REPORT.docx, matching the PDF.

Arial 11 body, Consolas for terminal excerpts, the same cover page and the same
table handling. The markdown parser is imported from build-pdf.py so the two
outputs cannot drift apart in what they consider a heading, a table or a code
block.

Word builds its own table of contents field, which needs a click to populate in
some viewers, so a static contents list is written instead. Page numbers go in
the footer as a field, because a .docx has no fixed pagination until it is laid
out by the reader.

Usage:  python scripts/build-docx.py [--md docs/REPORT.md] [--out REPORT.docx]
"""
import argparse
import importlib.util
import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

HERE = os.path.dirname(os.path.abspath(__file__))
spec = importlib.util.spec_from_file_location("buildpdf", os.path.join(HERE, "build-pdf.py"))
buildpdf = importlib.util.module_from_spec(spec)
spec.loader.exec_module(buildpdf)
parse, inline_segments, strip_inline, preprocess = (
    buildpdf.parse, buildpdf.inline_segments, buildpdf.strip_inline,
    buildpdf.preprocess)

TITLE = buildpdf.TITLE
SUBTITLE = buildpdf.SUBTITLE
AUTHOR = buildpdf.AUTHOR

BODY_PT = 11
CODE_PT = 8.5
TABLE_PT = 8.5


def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def add_page_number_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for instr, kind in (("begin", "fldCharType"), (None, "instrText"), ("end", "fldCharType")):
        if kind == "fldCharType":
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), instr)
        else:
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = " PAGE "
        run._r.append(el)
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6E, 0x6E, 0x6E)


def write_runs(par, text, base_size=BODY_PT):
    for seg_text, style, mono in inline_segments(text):
        run = par.add_run(seg_text)
        run.font.name = "Consolas" if mono else "Arial"
        # east-asian name must be set too or Word substitutes
        run._element.rPr.rFonts.set(qn("w:eastAsia"), run.font.name)
        run.font.size = Pt(base_size - 1.5 if mono else base_size)
        if "B" in style:
            run.bold = True
        if "I" in style:
            run.italic = True
        if mono:
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def build(md_path, out_path):
    md = preprocess(open(md_path, encoding="utf-8").read())
    blocks = parse(md)
    blocks = [b for b in blocks if not (b[0] == "h" and strip_inline(b[1]["text"]).lower()
                                        in ("table of contents", "contents"))]

    doc = Document()

    st = doc.styles["Normal"]
    st.font.name = "Arial"
    st.font.size = Pt(BODY_PT)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.15

    for s in doc.sections:
        s.page_width, s.page_height = Cm(21.0), Cm(29.7)
        s.left_margin = s.right_margin = Cm(2.0)
        s.top_margin = s.bottom_margin = Cm(1.8)

    # ------------------------------------------------------------ cover
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE); r.bold = True; r.font.size = Pt(26); r.font.name = "Arial"
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SUBTITLE); r.font.size = Pt(13); r.font.name = "Arial"
    r.font.color.rgb = RGBColor(0x46, 0x46, 0x46)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(AUTHOR); r.font.size = Pt(12); r.font.name = "Arial"
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(buildpdf.DATED)
    r.font.size = Pt(10); r.font.name = "Arial"
    r.font.color.rgb = RGBColor(0x6E, 0x6E, 0x6E)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(buildpdf.REPO)
    r.font.size = Pt(10); r.font.name = "Arial"
    r.font.color.rgb = RGBColor(0x6E, 0x6E, 0x6E)
    doc.paragraphs[-1].runs[-1].add_break(WD_BREAK.PAGE)

    # ---------------------------------------------------------- contents
    p = doc.add_paragraph()
    r = p.add_run("Contents"); r.bold = True; r.font.size = Pt(16); r.font.name = "Arial"
    for kind, payload in blocks:
        if kind == "h" and payload["level"] in (2, 3):
            label = strip_inline(payload["text"])
            par = doc.add_paragraph()
            par.paragraph_format.space_after = Pt(1)
            par.paragraph_format.left_indent = Cm(0.0 if payload["level"] == 2 else 0.7)
            rr = par.add_run(label)
            rr.font.name = "Arial"
            rr.font.size = Pt(10.5 if payload["level"] == 2 else 10)
            rr.bold = payload["level"] == 2
    doc.paragraphs[-1].runs[-1].add_break(WD_BREAK.PAGE)

    add_page_number_footer(doc.sections[0])

    # ------------------------------------------------------------ body
    sizes = {2: 15.5, 3: 12.5, 4: 11.5}
    for kind, payload in blocks:
        if kind == "h":
            lvl = payload["level"]
            par = doc.add_paragraph()
            par.paragraph_format.space_before = Pt(14 if lvl == 2 else 9)
            par.paragraph_format.space_after = Pt(4)
            if lvl == 2 and len(doc.paragraphs) > 3:
                par.paragraph_format.page_break_before = True
            r = par.add_run(strip_inline(payload["text"]))
            r.bold = True
            r.font.size = Pt(sizes.get(lvl, 11))
            r.font.name = "Arial"
        elif kind == "p":
            par = doc.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            write_runs(par, payload)
        elif kind == "li":
            par = doc.add_paragraph(style="List Bullet")
            par.paragraph_format.left_indent = Cm(0.7 + payload["indent"] * 0.25)
            par.paragraph_format.space_after = Pt(2)
            write_runs(par, payload["text"])
        elif kind == "quote":
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Cm(0.6)
            write_runs(par, payload)
            for r in par.runs:
                r.font.color.rgb = RGBColor(0x3C, 0x3C, 0x3C)
        elif kind == "code":
            lines = payload["lines"]
            if not lines:
                continue
            tbl = doc.add_table(rows=1, cols=1)
            tbl.style = "Table Grid"
            cell = tbl.rows[0].cells[0]
            set_cell_bg(cell, "F6F6F6")
            cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
            for l in lines:
                cp = cell.add_paragraph()
                cp.paragraph_format.space_after = Pt(0)
                cp.paragraph_format.space_before = Pt(0)
                cp.paragraph_format.line_spacing = 1.0
                rr = cp.add_run(l.replace("\t", "    "))
                rr.font.name = "Consolas"
                rr._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                rr.font.size = Pt(CODE_PT if max(len(x) for x in lines) <= 95 else 7.0)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        elif kind == "table":
            rows = payload
            ncols = max(len(r) for r in rows)
            rows = [r + [""] * (ncols - len(r)) for r in rows]
            tbl = doc.add_table(rows=len(rows), cols=ncols)
            tbl.style = "Table Grid"
            tbl.autofit = True
            for ri, row in enumerate(rows):
                for ci, txt in enumerate(row):
                    cell = tbl.cell(ri, ci)
                    cell.text = ""
                    par = cell.paragraphs[0]
                    par.paragraph_format.space_after = Pt(1)
                    par.paragraph_format.space_before = Pt(1)
                    write_runs(par, txt, base_size=TABLE_PT)
                    if ri == 0:
                        set_cell_bg(cell, "EEEEEE")
                        for r in par.runs:
                            r.bold = True
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        elif kind == "hr":
            par = doc.add_paragraph()
            par.paragraph_format.space_after = Pt(2)

    cp = doc.core_properties
    cp.title = TITLE
    cp.author = AUTHOR
    cp.subject = ""
    cp.keywords = ""
    cp.comments = ""
    cp.category = ""
    cp.last_modified_by = ""

    doc.save(out_path)
    print(f"wrote {out_path}")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--md", default="docs/REPORT.md")
    ap.add_argument("--out", default="REPORT.docx")
    a = ap.parse_args()
    build(a.md, a.out)


if __name__ == "__main__":
    main()
